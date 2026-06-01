#!/usr/bin/env python3
"""Build a PDF + DOCX cover letter from a markdown file.

PDF  -> headless Chrome (Playwright).
DOCX -> python-docx, single-column, no tables -> maximal ATS parseability.

Run:
    python3 build_cover_letter.py MyLetter.md            # outputs MyLetter.pdf + MyLetter.docx
    python3 build_cover_letter.py MyLetter.md --preview  # also writes MyLetter-preview.png
"""

import argparse
import re
from pathlib import Path

INLINE_TOKEN = re.compile(r'\[([^\]]+)\]\(([^)]+)\)|\*\*(.+?)\*\*')


def parse_inline(text):
    runs = []
    pos = 0
    for m in INLINE_TOKEN.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, None))
        if m.group(1) is not None:
            runs.append((m.group(1), False, m.group(2)))
        else:
            runs.append((m.group(3), True, None))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, None))
    return runs


def parse_letter(text):
    """Parse a cover letter markdown into structured fields.

    Expected format:
        Sender Name
        Contact line
        Date

        Recipient line 1
        Recipient line 2
        Recipient line 3

        ---

        Body paragraphs...
    """
    lines = text.splitlines()
    model = {
        "sender_name": "",
        "sender_contact": [],
        "date": "",
        "recipient": [],
        "body": [],
    }

    i = 0
    # sender name
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines):
        model["sender_name"] = lines[i].strip()
        i += 1

    # contact
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines):
        model["sender_contact"] = parse_inline(lines[i].strip())
        i += 1

    # date
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines):
        model["date"] = lines[i].strip()
        i += 1

    # recipient block: non-empty lines until --- or blank line followed by ---
    while i < len(lines) and not lines[i].strip():
        i += 1
    while i < len(lines):
        s = lines[i].strip()
        if re.match(r'^---+$', s):
            i += 1
            break
        if s:
            model["recipient"].append(s)
        i += 1

    # body: collect paragraphs (split on blank lines)
    # Short consecutive lines (closing, salutation) are joined with <br>; long lines are joined with space.
    current = []
    while i < len(lines):
        s = lines[i].strip()
        if re.match(r'^---+$', s):
            i += 1
            continue
        if s:
            current.append(s)
        else:
            if current:
                if len(current) > 1 and all(len(l) < 60 for l in current):
                    model["body"].append("<br>".join(current))
                else:
                    model["body"].append(" ".join(current))
                current = []
        i += 1
    if current:
        if len(current) > 1 and all(len(l) < 60 for l in current):
            model["body"].append("<br>".join(current))
        else:
            model["body"].append(" ".join(current))

    return model


CSS = """
* { box-sizing: border-box; margin: 0; padding: 0; }
@page { size: letter; margin: 0.9in 1in; }
body {
    font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
    font-size: 10.5pt;
    line-height: 1.4;
    color: #1a1a1a;
}
.sender-name {
    font-size: 14pt;
    font-weight: 700;
    margin-bottom: 2pt;
}
.sender-contact {
    font-size: 9.5pt;
    color: #444;
    margin-bottom: 18pt;
}
.sender-contact a { color: #1a1a1a; text-decoration: none; }
.date {
    margin-bottom: 14pt;
}
.recipient {
    margin-bottom: 18pt;
}
.recipient div {
    line-height: 1.6;
}
hr {
    border: none;
    border-top: 1px solid #bbb;
    margin: 18pt 0;
}
.body p {
    margin-bottom: 9pt;
    text-align: left;
}
"""


def runs_to_html(runs):
    out = []
    for txt, bold, href in runs:
        t = txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if href:
            out.append(f'<a href="{href}">{t}</a>')
        elif bold:
            out.append(f"<strong>{t}</strong>")
        else:
            out.append(t)
    return "".join(out)


def build_html(model):
    parts = []
    parts.append(f'<div class="sender-name">{model["sender_name"]}</div>')
    parts.append(f'<div class="sender-contact">{runs_to_html(model["sender_contact"])}</div>')
    parts.append(f'<div class="date">{model["date"]}</div>')

    if model["recipient"]:
        parts.append('<div class="recipient">')
        for line in model["recipient"]:
            parts.append(f'<div>{line}</div>')
        parts.append('</div>')

    parts.append('<hr>')
    parts.append('<div class="body">')
    for para in model["body"]:
        parts.append(f'<p>{para}</p>')
    parts.append('</div>')

    body = "\n".join(parts)
    return (f'<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">'
            f'<style>{CSS}</style></head><body>{body}</body></html>')


def render_pdf(html, out_pdf):
    from playwright.sync_api import sync_playwright
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page()
        page.set_content(html, wait_until="networkidle")
        page.pdf(path=str(out_pdf), format="Letter", print_background=True)
        browser.close()
    print(f"PDF  -> {out_pdf}")


def render_preview(html, out_png):
    from playwright.sync_api import sync_playwright
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page()
        page.set_content(html, wait_until="networkidle")
        page.set_viewport_size({"width": int((8.5 - 2.0) * 96), "height": 1300})
        page.screenshot(path=str(out_png), full_page=True)
        height = page.evaluate("document.body.scrollHeight")
        browser.close()
    pages = height / ((11 - 1.8) * 96)
    print(f"PNG  -> {out_png}  (~{pages:.2f} page(s) — READ this file to validate)")


# ──────────────────────────────────────────────────────────────────────────
# DOCX RENDERER (python-docx) — ATS-clean, single column, no tables
# ──────────────────────────────────────────────────────────────────────────
def render_docx(model, out_docx):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT = "Calibri"
    DARK = RGBColor(0x1A, 0x1A, 0x1A)
    GRAY = RGBColor(0x44, 0x44, 0x44)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.9)
        sec.left_margin = sec.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    def add_hyperlink(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color"); color.set(qn("w:val"), "1A1A1A")
        rpr.append(color)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "none"); rpr.append(u)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        link.append(run)
        paragraph._p.append(link)

    def emit_runs(paragraph, runs, size=10.5, color=DARK):
        for txt, bold, href in runs:
            if href:
                add_hyperlink(paragraph, href, txt)
            else:
                r = paragraph.add_run(txt)
                r.bold = bold
                r.font.size = Pt(size)
                r.font.color.rgb = color

    # ── Sender name
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(model["sender_name"]); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = DARK

    # ── Sender contact (with hyperlinks)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    emit_runs(p, model["sender_contact"], size=9.5, color=GRAY)

    # ── Date
    if model["date"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(model["date"]); r.font.size = Pt(10.5); r.font.color.rgb = DARK

    # ── Recipient block
    for line in model["recipient"]:
        p = doc.add_paragraph()
        r = p.add_run(line); r.font.size = Pt(10.5); r.font.color.rgb = DARK
    if model["recipient"]:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(14)

    # ── Body paragraphs (split <br>-joined closing lines into line breaks)
    for para in model["body"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(9)
        for idx, seg in enumerate(para.split("<br>")):
            if idx > 0:
                p.add_run().add_break()
            r = p.add_run(seg); r.font.size = Pt(10.5); r.font.color.rgb = DARK

    doc.save(str(out_docx))
    print(f"DOCX -> {out_docx}")


def main():
    ap = argparse.ArgumentParser(description="Build a PDF cover letter from markdown.")
    ap.add_argument("input", help="markdown cover letter file")
    ap.add_argument("--out", help="output basename (default: derived from input filename)")
    ap.add_argument("--preview", action="store_true",
                    help="also write <base>-preview.png for visual validation")
    args = ap.parse_args()

    in_md = Path(args.input)
    if not in_md.is_file():
        ap.error(f"input file not found: {in_md}")

    base = args.out if args.out else in_md.stem
    out_pdf = in_md.parent / f"{base}.pdf"
    out_docx = in_md.parent / f"{base}.docx"

    model = parse_letter(in_md.read_text(encoding="utf-8"))
    html = build_html(model)
    render_pdf(html, out_pdf)
    render_docx(model, out_docx)
    if args.preview:
        render_preview(html, in_md.parent / f"{base}-preview.png")


if __name__ == "__main__":
    main()
