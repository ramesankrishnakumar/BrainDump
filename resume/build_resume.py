#!/usr/bin/env python3
"""Build ATS-friendly PDF + DOCX from a markdown resume.

PDF  -> headless Chrome (Playwright), pixel-styled like the reference.
DOCX -> python-docx, single-column, no tables -> maximal ATS parseability.

Run:
    python3 build_resume.py                      # uses last-resume-revised-4.md
    python3 build_resume.py my-resume.md         # outputs my-resume.pdf + my-resume.docx
    python3 build_resume.py in.md --out Resume   # outputs Resume.pdf + Resume.docx

Both files are generated every run.
"""

import argparse
import re
from pathlib import Path

HERE = Path(__file__).parent
DEFAULT_MD = HERE / "last-resume-revised-4.md"


# ──────────────────────────────────────────────────────────────────────────
# 1. PARSE MARKDOWN -> STRUCTURED MODEL
# ──────────────────────────────────────────────────────────────────────────
# Inline runs: list of (text, bold, href). href is None for plain/bold text.

INLINE_TOKEN = re.compile(r'\[([^\]]+)\]\(([^)]+)\)|\*\*(.+?)\*\*')


def parse_inline(text):
    """Return a list of (text, bold, href) run tuples."""
    runs = []
    pos = 0
    for m in INLINE_TOKEN.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, None))
        if m.group(1) is not None:          # link [label](url)
            runs.append((m.group(1), False, m.group(2)))
        else:                               # **bold**
            runs.append((m.group(3), True, None))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, None))
    return runs


def split_role(raw):
    """'Company — Role | Aug 2022 – Present' -> (left, right)."""
    m = re.search(r'\|\s*([A-Za-z]+\s+\d{4}.*)$', raw)
    if m:
        return raw[:m.start()].strip().rstrip('|').strip(), m.group(1).strip()
    return raw, ""


def parse_markdown(text):
    """Return dict: name, contact(runs), tagline(runs), sections[]."""
    lines = text.split("\n")
    model = {"name": "", "contact": [], "tagline": [], "sections": []}
    state = "name"
    section = None
    role = None

    for line in lines:
        s = line.strip()
        if not s:
            continue

        if state == "name":
            model["name"] = s
            state = "contact"
            continue
        if state == "contact":
            model["contact"] = parse_inline(s)
            state = "tagline"
            continue
        if state == "tagline":
            if re.match(r'^\*\*.+\*\*$', s):
                model["tagline"] = parse_inline(s[2:-2])
                state = "body"
                continue
            if re.match(r'^---+$', s):
                state = "body"
                continue
            # fall through if neither (defensive)

        if re.match(r'^---+$', s):
            role = None
            continue

        if s.startswith("## "):
            section = {"heading": s[3:].strip(), "blocks": []}
            model["sections"].append(section)
            role = None
            continue

        if s.startswith("### "):
            left, right = split_role(s[4:])
            role = {"type": "role", "left": parse_inline(left),
                    "right": right, "bullets": []}
            section["blocks"].append(role)
            continue

        if s.startswith("- "):
            if role is None:                # bullet without a role header
                role = {"type": "loose_bullets", "bullets": []}
                section["blocks"].append(role)
            role["bullets"].append(parse_inline(s[2:]))
            continue

        # plain paragraph (PROFILE / EDUCATION / SKILLS); strip trailing MD break
        section["blocks"].append({"type": "para", "runs": parse_inline(s.rstrip())})
        role = None

    return model


# ──────────────────────────────────────────────────────────────────────────
# 2. PDF RENDERER (HTML -> Chrome)
# ──────────────────────────────────────────────────────────────────────────
CSS = """
* { box-sizing: border-box; margin: 0; padding: 0; }
@page { size: letter; margin: 0.5in 0.62in; }
body { font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
       font-size: 9.8pt; line-height: 1.32; color: #1a1a1a; }
.name { font-size: 22pt; font-weight: 700; text-align: center;
        letter-spacing: 0.5px; margin-bottom: 4pt; }
.contact { text-align: center; font-size: 8.8pt; color: #333; margin-bottom: 3pt; }
.contact a { color: #1a1a1a; text-decoration: none; }
.tagline { text-align: center; font-size: 8.8pt; font-weight: 700; color: #1a1a1a;
           letter-spacing: 0.4px; margin-bottom: 8pt; }
.header-hr { border: none; border-top: 1.8px solid #1a1a1a; }
h2 { font-size: 9pt; font-weight: 700; text-transform: uppercase; letter-spacing: 1.6px;
     border-bottom: 0.75px solid #999; padding-bottom: 2.5pt;
     margin-top: 8pt; margin-bottom: 4pt; page-break-after: avoid; }
.role-row { display: flex; justify-content: space-between; align-items: baseline;
            gap: 12pt; margin-top: 5pt; margin-bottom: 1pt; page-break-after: avoid; }
.role-left { font-weight: 700; font-size: 9.8pt; }
.role-right { font-weight: 400; font-size: 9pt; color: #444; white-space: nowrap; }
ul { padding-left: 13pt; margin: 1pt 0 2pt; }
li { margin-bottom: 2.5pt; line-height: 1.3; page-break-inside: avoid; }
p { margin-bottom: 1.5pt; }
strong { font-weight: 700; }
a { color: #1a1a1a; text-decoration: none; }
"""


def runs_to_html(runs):
    out = []
    for txt, bold, href in runs:
        t = txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if href:
            out.append(f'<a href="{href}">{t}</a>')
        elif bold:
            out.append(f"<strong>{t}</strong>")
        else:
            out.append(t)
    return "".join(out)


def build_html(model):
    p = [f'<div class="name">{model["name"]}</div>',
         f'<p class="contact">{runs_to_html(model["contact"])}</p>',
         f'<p class="tagline">{runs_to_html(model["tagline"])}</p>',
         '<hr class="header-hr">']
    for sec in model["sections"]:
        p.append(f'<h2>{sec["heading"]}</h2>')
        for blk in sec["blocks"]:
            if blk["type"] == "para":
                p.append(f'<p>{runs_to_html(blk["runs"])}</p>')
            elif blk["type"] in ("role", "loose_bullets"):
                if blk["type"] == "role":
                    p.append(
                        '<div class="role-row">'
                        f'<span class="role-left">{runs_to_html(blk["left"])}</span>'
                        f'<span class="role-right">{blk["right"]}</span></div>')
                if blk["bullets"]:
                    p.append("<ul>")
                    p += [f'<li>{runs_to_html(b)}</li>' for b in blk["bullets"]]
                    p.append("</ul>")
    body = "\n".join(p)
    title = model["name"] or "Resume"
    return (f'<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">'
            f'<title>{title} — Resume</title>'
            f'<style>{CSS}</style></head><body>{body}</body></html>')


def render_pdf(model, out_pdf):
    from playwright.sync_api import sync_playwright
    html = build_html(model)
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page()
        page.set_content(html, wait_until="networkidle")
        page.pdf(path=str(out_pdf), format="Letter", print_background=True,
                 margin={"top": "0.5in", "bottom": "0.5in",
                         "left": "0.65in", "right": "0.65in"})
        browser.close()
    print(f"PDF  -> {out_pdf}")


# ──────────────────────────────────────────────────────────────────────────
# 2b. PREVIEW RENDERER (HTML -> PNG)  ── AGENT VALIDATION: MUST DO ──
# ──────────────────────────────────────────────────────────────────────────
# Chrome cannot render a .pdf inline (it triggers a download), so you CANNOT
# screenshot the PDF directly. To visually validate the output, render the same
# HTML to a full-page PNG and READ it. An agent editing the layout/CSS MUST:
#   1. run `python3 build_resume.py --preview`
#   2. open the generated *-preview.png with the Read tool and confirm:
#      centered name, section rules, dates right-aligned, clean page break,
#      no overflow/clipping — compare against the reference image.
# Viewport width mimics the Letter printable area (8.5in - 1.24in margins @96dpi).
def render_preview(model, out_png):
    from playwright.sync_api import sync_playwright
    html = build_html(model)
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page()
        page.set_content(html, wait_until="networkidle")
        page.set_viewport_size({"width": int((8.5 - 1.24) * 96), "height": 1300})
        page.screenshot(path=str(out_png), full_page=True)
        height = page.evaluate("document.body.scrollHeight")
        browser.close()
    pages = height / ((11 - 1.0) * 96)  # content height / usable Letter page height
    print(f"PNG  -> {out_png}  (~{pages:.2f} page(s) — READ this file to validate)")


# ──────────────────────────────────────────────────────────────────────────
# 3. DOCX RENDERER (python-docx) — ATS-clean, single column, no tables
# ──────────────────────────────────────────────────────────────────────────
def render_docx(model, out_docx):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT = "Calibri"
    GRAY = RGBColor(0x55, 0x55, 0x55)
    DARK = RGBColor(0x1A, 0x1A, 0x1A)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.5)
        sec.left_margin = sec.right_margin = Inches(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.08

    def add_hyperlink(paragraph, url, text, size=10):
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color"); color.set(qn("w:val"), "1A1A1A")
        rpr.append(color)
        if size:
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2))
            rpr.append(sz)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "none"); rpr.append(u)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        link.append(run)
        paragraph._p.append(link)

    def emit_runs(paragraph, runs, size=10, color=DARK, plain=False):
        # plain=True renders linked text as a normal <w:r> run instead of a
        # <w:hyperlink>. Resume parsers that iterate only top-level runs skip
        # text nested inside <w:hyperlink>, so contact email/URLs must be plain.
        for txt, bold, href in runs:
            if href and not plain:
                add_hyperlink(paragraph, href, txt, size=size)
            else:
                r = paragraph.add_run(txt)
                r.bold = bold
                r.font.size = Pt(size)
                r.font.color.rgb = color

    def bottom_border(paragraph):
        ppr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "999999")
        pbdr.append(bottom)
        ppr.append(pbdr)

    # ── Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(model["name"]); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK
    # No outlineLvl on the name: parsers that read outline-0 as a document title
    # (rather than body content) can fail to extract it as the candidate name.

    # ── Contact
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    emit_runs(p, model["contact"], size=9, color=RGBColor(0x33, 0x33, 0x33), plain=True)

    # ── Tagline
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    for txt, _bold, _href in model["tagline"]:
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = DARK

    page_w = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    for sec in model["sections"]:
        # ── Section heading + rule
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(3)
        r = h.add_run(sec["heading"].upper())
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = DARK
        rpr = r._element.get_or_add_rPr()
        spc = OxmlElement("w:spacing"); spc.set(qn("w:val"), "30"); rpr.append(spc)
        ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), "1")
        h._p.get_or_add_pPr().append(ol)
        bottom_border(h)

        for blk in sec["blocks"]:
            if blk["type"] == "para":
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                emit_runs(p, blk["runs"])
            elif blk["type"] in ("role", "loose_bullets"):
                if blk["type"] == "role":
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(5)
                    p.paragraph_format.space_after = Pt(1)
                    tab_stops = p.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(page_w, WD_TAB_ALIGNMENT.RIGHT)
                    for txt, _b, href in blk["left"]:
                        rr = p.add_run(txt); rr.bold = True; rr.font.size = Pt(10)
                    if blk["right"]:
                        rr = p.add_run(" \t" + blk["right"])
                        rr.font.size = Pt(10); rr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                for bullet in blk["bullets"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.space_after = Pt(2)
                    emit_runs(p, bullet)

    doc.save(str(out_docx))
    print(f"DOCX -> {out_docx}")


def main():
    ap = argparse.ArgumentParser(description="Build ATS-friendly PDF + DOCX from a markdown resume.")
    ap.add_argument("input", nargs="?", default=str(DEFAULT_MD),
                    help="markdown resume file (default: last-resume-revised-4.md)")
    ap.add_argument("--out", help="output basename (default: derived from input filename)")
    ap.add_argument("--preview", action="store_true",
                    help="also write <base>-preview.png for visual validation (see render_preview)")
    args = ap.parse_args()

    in_md = Path(args.input)
    if not in_md.is_file():
        ap.error(f"input file not found: {in_md}")

    base = args.out if args.out else in_md.stem
    out_pdf = in_md.parent / f"{base}.pdf"
    out_docx = in_md.parent / f"{base}.docx"

    model = parse_markdown(in_md.read_text(encoding="utf-8"))
    render_pdf(model, out_pdf)
    render_docx(model, out_docx)
    if args.preview:
        render_preview(model, in_md.parent / f"{base}-preview.png")


if __name__ == "__main__":
    main()
